"""
Resume Analyzer Module
-----------------------
Provides:
  - LOADED_RESUME: module-level cache of the last extracted resume text
  - extract_text_from_file(filepath): extracts plain text from PDF/DOC/DOCX
  - analyze_resume(chat, query): sends the resume text to the AI chat handler
    and returns the analysis as a string
"""

import os
import logging

logger = logging.getLogger(__name__)

# Holds the most recently uploaded & extracted resume text.
# Set by app.py's /api/upload route after a successful extraction.
LOADED_RESUME = None


def extract_text_from_file(filepath: str) -> str:
    """
    Extract plain text from a PDF, DOC, or DOCX file.

    Raises:
        ValueError: if the file type is unsupported or no text could be extracted.
        Exception: re-raises underlying library errors so the caller can log them.
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = filepath.rsplit('.', 1)[-1].lower()

    if ext == 'pdf':
        text = _extract_from_pdf(filepath)
    elif ext == 'docx':
        text = _extract_from_docx(filepath)
    elif ext == 'doc':
        # Legacy .doc (binary format) - try docx reader first as a fallback,
        # then fall back to a raw-text best-effort read.
        try:
            text = _extract_from_docx(filepath)
        except Exception:
            text = _extract_raw_text(filepath)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")

    text = text.strip()
    if not text:
        raise ValueError(
            "No readable text found in the file. "
            "It may be a scanned/image-based document that requires OCR."
        )

    logger.info(f"Extracted {len(text)} characters from {filepath}")
    return text


def _extract_from_pdf(filepath: str) -> str:
    import PyPDF2

    text_parts = []
    with open(filepath, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)

    return "\n".join(text_parts)


def _extract_from_docx(filepath: str) -> str:
    from docx import Document

    doc = Document(filepath)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Also pull text out of tables (common in resumes)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    return "\n".join(parts)


def _extract_raw_text(filepath: str) -> str:
    """Best-effort fallback for legacy .doc files without python-docx support."""
    with open(filepath, 'rb') as f:
        raw = f.read()

    # Strip non-printable bytes, keep readable ASCII chunks
    text = raw.decode('latin-1', errors='ignore')
    printable = ''.join(ch if (32 <= ord(ch) < 127 or ch in '\n\r\t') else ' ' for ch in text)
    # Collapse excessive whitespace
    lines = [line.strip() for line in printable.splitlines()]
    lines = [line for line in lines if len(line) > 2]
    return "\n".join(lines)


def analyze_resume(chat, query: str) -> str:
    """
    Send the resume content (and any extra user instructions) to the AI
    and return a structured analysis as plain text.

    Args:
        chat: object with a send_message(prompt) method returning an object
              that has a `.text` attribute (matches ResponseHandler in app.py).
        query: the resume text (and/or extra instructions appended by the user).

    Returns:
        The AI's analysis as a string.
    """
    if not query or not query.strip():
        return "No resume content provided. Please upload a resume file (PDF/DOC/DOCX) or paste resume text."

    prompt = f"""You are an expert resume reviewer and career coach. Analyze the following resume and provide a detailed, structured assessment.

Cover these areas:
1. **Overall Impression** - a brief summary of the candidate's profile
2. **Strengths** - what stands out positively (skills, projects, achievements)
3. **Areas for Improvement** - formatting, missing sections, weak phrasing, etc.
4. **ATS Compatibility** - keyword usage, formatting issues that could hurt ATS parsing
5. **Skills Gap** - skills commonly expected for this candidate's target roles that are missing
6. **Actionable Suggestions** - specific, prioritized changes to make the resume stronger

Resume content:
\"\"\"
{query}
\"\"\"

Provide the analysis in clear sections using the headings above."""

    try:
        response = chat.send_message(prompt)
        result = getattr(response, 'text', '') or ''
        if not result.strip():
            return "The AI did not return any analysis. Please try again."
        return result
    except Exception as e:
        logger.error(f"Resume analysis failed: {e}", exc_info=True)
        raise